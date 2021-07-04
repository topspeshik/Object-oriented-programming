from docxtpl import DocxTemplate
from interface import *
from docx import Document
from docx.shared import Inches

if __name__ == "__main__":
    import sys
    app = QtWidgets.QApplication(sys.argv)
    TabWidget = QtWidgets.QTabWidget()
    ui = Ui_TabWidget()
    ui.setupUi(TabWidget)
    TabWidget.show()


    def create_word():

        document = Document("gg.docx")
        arrWords.append(ui.p1_document.text()+".docx")
        document.save(arrWords[arrI[0]])
        ui.p1_document.setText('')

    def table1():

        document = Document(arrWords[arrI[0]])

        table = document.tables[3]

        if (ui.t1_compet.toPlainText() == "") and (ui.t1_indicator.toPlainText() == ""):
            table.add_row()
            table.rows[len(table.rows) - 1].cells[0].merge((table.rows[len(table.rows) - 2].cells[0]))
            table.rows[len(table.rows) - 1].cells[1].merge((table.rows[len(table.rows) - 2].cells[1]))
            table.rows[len(table.rows) - 1].cells[2].text = ui.t1_result.toPlainText()

        elif (ui.t1_compet.toPlainText() == "") and (ui.t1_result.toPlainText() == ""):
            table.add_row()
            table.rows[len(table.rows) - 1].cells[0].merge((table.rows[len(table.rows) - 2].cells[0]))
            table.rows[len(table.rows) - 1].cells[1].text = ui.t1_indicator.toPlainText()
            table.rows[len(table.rows) - 1].cells[2].merge((table.rows[len(table.rows) - 2].cells[2]))
        elif (ui.t1_indicator.toPlainText() == "") and (ui.t1_result.toPlainText() == ""):
            table.add_row()
            table.rows[len(table.rows) - 1].cells[0].text = ui.t1_compet.toPlainText()
            table.rows[len(table.rows) - 1].cells[1].merge((table.rows[len(table.rows) - 2].cells[1]))
            table.rows[len(table.rows) - 1].cells[2].merge((table.rows[len(table.rows) - 2].cells[2]))
        elif ui.t1_compet.toPlainText() == "":
            table.add_row()
            table.rows[len(table.rows) - 1].cells[0].merge((table.rows[len(table.rows) - 2].cells[0]))
            table.rows[len(table.rows) - 1].cells[1].text = ui.t1_indicator.toPlainText()
            table.rows[len(table.rows) - 1].cells[2].text = ui.t1_result.toPlainText()
        elif ui.t1_indicator.toPlainText() == "":
            table.add_row()
            table.rows[len(table.rows) - 1].cells[0].text = ui.t1_compet.toPlainText()
            table.rows[len(table.rows) - 1].cells[1].merge((table.rows[len(table.rows) - 2].cells[1]))
            table.rows[len(table.rows) - 1].cells[2].text = ui.t1_result.toPlainText()
        elif ui.t1_result.toPlainText() == "":
            table.add_row()
            table.rows[len(table.rows) - 1].cells[0].text = ui.t1_compet.toPlainText()
            table.rows[len(table.rows) - 1].cells[1].text = ui.t1_indicator.toPlainText()
            table.rows[len(table.rows) - 1].cells[2].merge((table.rows[len(table.rows) - 2].cells[2]))
        else:
            table.add_row()
            table.rows[len(table.rows) - 1].cells[0].text = ui.t1_compet.toPlainText()
            table.rows[len(table.rows) - 1].cells[1].text = ui.t1_indicator.toPlainText()
            table.rows[len(table.rows) - 1].cells[2].text = ui.t1_result.toPlainText()

        document.save(arrWords[arrI[0]])

        ui.t1_compet.setPlainText('')
        ui.t1_indicator.setPlainText('')
        ui.t1_result.setPlainText('')

    def is_int(s):
        try:
            int(s)
            return True
        except ValueError:
            return False

    def table2():
        document = Document(arrWords[arrI[0]])

        table = document.tables[4]
        table.add_row()

        if is_int(ui.t2_razdel_2.toPlainText()):
            table.rows[len(table.rows) - 1].cells[0].add_paragraph(ui.t2_razdel_2.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            arrNumber.append(float(ui.t2_razdel_2.toPlainText()))
            p = table.rows[len(table.rows) - 1].cells[1].add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0)
            p.add_run(ui.t2_razdel.toPlainText()).bold = True
        elif ui.t2_razdel.toPlainText() == 'семестр':
            table.rows[len(table.rows) - 1].cells[0].add_paragraph(ui.t2_razdel_2.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            p = table.rows[len(table.rows) - 1].cells[1].add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0)
            p.add_run(ui.t2_razdel.toPlainText()).bold = True
        elif ui.t2_razdel.toPlainText() == 'Семестр':
            table.rows[len(table.rows) - 1].cells[0].add_paragraph(ui.t2_razdel_2.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            p = table.rows[len(table.rows) - 1].cells[1].add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0)
            p.add_run(ui.t2_razdel.toPlainText()).bold = True
        else:
            arrNumber[len(arrNumber)-1] += 0.1
            table.rows[len(table.rows) - 1].cells[0].add_paragraph(str( round(arrNumber[len(arrNumber)-1], 1))).paragraph_format.first_line_indent = Inches(0)
            table.rows[len(table.rows) - 1].cells[1].add_paragraph(ui.t2_razdel.toPlainText()).paragraph_format.first_line_indent = Inches(0)


        table.rows[len(table.rows) - 1].cells[2].add_paragraph(ui.t2_semestr.toPlainText()).paragraph_format.first_line_indent = Inches(0)
        table.rows[len(table.rows) - 1].cells[3].add_paragraph(ui.t2_lekci.toPlainText()).paragraph_format.first_line_indent = Inches(0)
        table.rows[len(table.rows) - 1].cells[4].add_paragraph(ui.t2_seminar.toPlainText()).paragraph_format.first_line_indent = Inches(0)
        table.rows[len(table.rows) - 1].cells[5].add_paragraph(ui.t2_consult.toPlainText()).paragraph_format.first_line_indent = Inches(0)
        table.rows[len(table.rows) - 1].cells[6].add_paragraph(ui.t1_compet_7.toPlainText()).paragraph_format.first_line_indent = Inches(0)
        table.rows[len(table.rows) - 1].cells[7].add_paragraph(ui.t1_compet_8.toPlainText()).paragraph_format.first_line_indent = Inches(0)
        if table.rows[len(table.rows) - 1].cells[0].text.strip() == "":
            arrSemestr.append(int(ui.t2_lekci.toPlainText()))
            arrSeminar.append(int(ui.t2_seminar.toPlainText()))
            arrSamr.append(int(ui.t1_compet_7.toPlainText()))

        document.save(arrWords[arrI[0]])
        ui.t2_razdel_2.setPlainText('')
        ui.t2_razdel.setPlainText('')
        ui.t2_semestr.setPlainText('')
        ui.t2_lekci.setPlainText('')
        ui.t2_seminar.setPlainText('')
        ui.t2_consult.setPlainText('')
        ui.t1_compet_7.setPlainText('')
        ui.t1_compet_8.setPlainText('')


    def table2_count():
        document = Document(arrWords[arrI[0]])
        table = document.tables[4]
        table.add_row()
        table.rows[len(table.rows) - 1].cells[1].add_paragraph("Итого часов").paragraph_format.first_line_indent = Inches(0)
        table.rows[len(table.rows) - 1].cells[3].add_paragraph(str(sum(arrSemestr))).paragraph_format.first_line_indent = Inches(0)
        table.rows[len(table.rows) - 1].cells[4].add_paragraph(str(sum(arrSeminar))).paragraph_format.first_line_indent = Inches(0)
        table.rows[len(table.rows) - 1].cells[6].add_paragraph(str(sum(arrSamr))).paragraph_format.first_line_indent = Inches(0)
        table.rows[len(table.rows) - 1].cells[7].add_paragraph(ui.pIV_hoursEx.text()).paragraph_format.first_line_indent = Inches(0)
        document.save(arrWords[arrI[0]])

    def table3():
        document = Document(arrWords[arrI[0]])

        table = document.tables[5]

        print(table.rows[0].cells[0].text.strip())
        t3_sumHours.append(int(ui.t3_zatrat.toPlainText()))
        if (ui.t3_vid.toPlainText() == "") and (ui.t3_ocenka.toPlainText() == ""):
            table.add_row()
            table.rows[len(table.rows) - 1].cells[2].merge((table.rows[len(table.rows) - 2].cells[2]))
            table.rows[len(table.rows) - 1].cells[5].merge((table.rows[len(table.rows) - 2].cells[5]))
            table.rows[len(table.rows) - 1].cells[1].add_paragraph(ui.t3_nazvanie.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            table.rows[len(table.rows) - 1].cells[0].add_paragraph(ui.t3_semestr.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            table.rows[len(table.rows) - 1].cells[3].add_paragraph(ui.t3_sroki.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            table.rows[len(table.rows) - 1].cells[4].add_paragraph(ui.t3_zatrat.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            table.rows[len(table.rows) - 1].cells[6].add_paragraph(ui.t3_uchebn.toPlainText()).paragraph_format.first_line_indent = Inches(0)

        elif (ui.t3_vid.toPlainText() == "") and (ui.t3_uchebn.toPlainText() == ""):
            table.add_row()
            table.rows[len(table.rows) - 1].cells[2].merge((table.rows[len(table.rows) - 2].cells[2]))
            table.rows[len(table.rows) - 1].cells[6].merge((table.rows[len(table.rows) - 2].cells[6]))
            table.rows[len(table.rows) - 1].cells[1].add_paragraph(ui.t3_nazvanie.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            table.rows[len(table.rows) - 1].cells[0].add_paragraph(ui.t3_semestr.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            table.rows[len(table.rows) - 1].cells[3].add_paragraph(ui.t3_sroki.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            table.rows[len(table.rows) - 1].cells[4].add_paragraph(ui.t3_zatrat.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            table.rows[len(table.rows) - 1].cells[5].add_paragraph(ui.t3_ocenka.toPlainText()).paragraph_format.first_line_indent = Inches(0)
        elif (ui.t3_uchebn.toPlainText() == "") and (ui.t3_ocenka.toPlainText() == ""):
            table.add_row()
            table.rows[len(table.rows) - 1].cells[5].merge((table.rows[len(table.rows) - 2].cells[5]))
            table.rows[len(table.rows) - 1].cells[6].merge((table.rows[len(table.rows) - 2].cells[6]))
            table.rows[len(table.rows) - 1].cells[1].add_paragraph(ui.t3_nazvanie.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            table.rows[len(table.rows) - 1].cells[0].add_paragraph(ui.t3_semestr.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            table.rows[len(table.rows) - 1].cells[3].add_paragraph(ui.t3_sroki.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            table.rows[len(table.rows) - 1].cells[4].add_paragraph(ui.t3_zatrat.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            table.rows[len(table.rows) - 1].cells[2].add_paragraph(ui.t3_vid.toPlainText()).paragraph_format.first_line_indent = Inches(0)
        elif ui.t3_uchebn.toPlainText() == "":
            table.add_row()
            table.rows[len(table.rows) - 1].cells[5].add_paragraph(ui.t3_ocenka.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            table.rows[len(table.rows) - 1].cells[6].merge((table.rows[len(table.rows) - 2].cells[6]))
            table.rows[len(table.rows) - 1].cells[1].add_paragraph(ui.t3_nazvanie.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            table.rows[len(table.rows) - 1].cells[0].add_paragraph(ui.t3_semestr.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            table.rows[len(table.rows) - 1].cells[3].add_paragraph(ui.t3_sroki.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            table.rows[len(table.rows) - 1].cells[4].add_paragraph(ui.t3_zatrat.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            table.rows[len(table.rows) - 1].cells[2].add_paragraph(ui.t3_vid.toPlainText()).paragraph_format.first_line_indent = Inches(0)
        elif ui.t3_ocenka.toPlainText() == "":
            table.add_row()
            table.rows[len(table.rows) - 1].cells[5].merge((table.rows[len(table.rows) - 2].cells[5]))
            table.rows[len(table.rows) - 1].cells[6].add_paragraph(ui.t3_uchebn.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            table.rows[len(table.rows) - 1].cells[1].add_paragraph(ui.t3_nazvanie.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            table.rows[len(table.rows) - 1].cells[0].add_paragraph(ui.t3_semestr.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            table.rows[len(table.rows) - 1].cells[3].add_paragraph(ui.t3_sroki.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            table.rows[len(table.rows) - 1].cells[4].add_paragraph(ui.t3_zatrat.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            table.rows[len(table.rows) - 1].cells[2].add_paragraph(ui.t3_vid.toPlainText()).paragraph_format.first_line_indent = Inches(0)
        elif ui.t3_vid.toPlainText() == "":
            table.add_row()
            table.rows[len(table.rows) - 1].cells[2].merge((table.rows[len(table.rows) - 2].cells[2]))
            table.rows[len(table.rows) - 1].cells[6].add_paragraph(ui.t3_uchebn.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            table.rows[len(table.rows) - 1].cells[1].add_paragraph(ui.t3_nazvanie.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            table.rows[len(table.rows) - 1].cells[0].add_paragraph(ui.t3_semestr.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            table.rows[len(table.rows) - 1].cells[3].add_paragraph(ui.t3_sroki.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            table.rows[len(table.rows) - 1].cells[4].add_paragraph(ui.t3_zatrat.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            table.rows[len(table.rows) - 1].cells[5].add_paragraph(ui.t3_ocenka.toPlainText()).paragraph_format.first_line_indent = Inches(0)
        else:
            table.add_row()
            table.rows[len(table.rows) - 1].cells[2].add_paragraph(ui.t3_vid.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            table.rows[len(table.rows) - 1].cells[6].add_paragraph(ui.t3_uchebn.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            table.rows[len(table.rows) - 1].cells[1].add_paragraph(ui.t3_nazvanie.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            table.rows[len(table.rows) - 1].cells[0].add_paragraph(ui.t3_semestr.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            table.rows[len(table.rows) - 1].cells[3].add_paragraph(ui.t3_sroki.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            table.rows[len(table.rows) - 1].cells[4].add_paragraph(ui.t3_zatrat.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            table.rows[len(table.rows) - 1].cells[5].add_paragraph(ui.t3_ocenka.toPlainText()).paragraph_format.first_line_indent = Inches(0)

        document.save(arrWords[arrI[0]])
        ui.t3_vid.setPlainText('')
        ui.t3_uchebn.setPlainText('')
        ui.t3_nazvanie.setPlainText('')
        ui.t3_semestr.setPlainText('')
        ui.t3_sroki.setPlainText('')
        ui.t3_zatrat.setPlainText('')
        ui.t3_ocenka.setPlainText('')


    def table3_count():
        document = Document(arrWords[arrI[0]])
        table = document.tables[5]
        table.add_row()
        table.rows[len(table.rows) - 1].cells[0].merge(table.rows[len(table.rows) - 1].cells[3])
        table.rows[len(table.rows) - 1].cells[0].add_paragraph("Общая трудоемкость самостоятельной работы по дисциплине (час) ").paragraph_format.first_line_indent = Inches(0)
        table.rows[len(table.rows) - 1].cells[4].add_paragraph(str(sum(t3_sumHours)/2)).paragraph_format.first_line_indent = Inches(0)
        table.add_row()
        table.rows[len(table.rows) - 1].cells[0].merge(table.rows[len(table.rows) - 1].cells[3])
        table.rows[len(table.rows) - 1].cells[0].add_paragraph("Из них объем самостоятельной работы  с использованием электронного обучения и дистанционных образовательных технологий (час) ").paragraph_format.first_line_indent = Inches(0)
        table.rows[len(table.rows) - 1].cells[4].add_paragraph(
            str(sum(t3_sumHours) / 2)).paragraph_format.first_line_indent = Inches(0)
        table.add_row()
        table.rows[len(table.rows) - 1].cells[0].merge(table.rows[len(table.rows) - 1].cells[3])
        table.rows[len(table.rows) - 1].cells[0].add_paragraph(
            "Бюджет времени самостоятельной работы, предусмотренный учебным планом для данной дисциплины (час) ").paragraph_format.first_line_indent = Inches(
            0)
        table.rows[len(table.rows) - 1].cells[4].add_paragraph(
            str(sum(t3_sumHours) / 2)).paragraph_format.first_line_indent = Inches(0)
        document.save(arrWords[arrI[0]])


    def table4():
        document = Document(arrWords[arrI[0]])

        table = document.tables[7]

        table.add_row()

        if is_int(ui.t4_nomer.toPlainText()):
            table.rows[len(table.rows) - 1].cells[1].add_paragraph(ui.t4_nomer.toPlainText()).paragraph_format.first_line_indent = Inches(0)
            arrNumber_t4.append(float(ui.t4_nomer.toPlainText()))
        else:

            arrNumber_t4[len(arrNumber_t4) - 1] += 0.1
            table.rows[len(table.rows) - 1].cells[1].add_paragraph(str(round(arrNumber_t4[len(arrNumber_t4) - 1], 1))).paragraph_format.first_line_indent = Inches(0)
        arrCount_t4[0] +=1
        table.rows[len(table.rows) - 1].cells[0].add_paragraph(str(arrCount_t4[0])).paragraph_format.first_line_indent = Inches(0)
        table.rows[len(table.rows) - 1].cells[2].add_paragraph(ui.t4_name.toPlainText()).paragraph_format.first_line_indent = Inches(0)
        table.rows[len(table.rows) - 1].cells[3].add_paragraph(ui.t4_trud.toPlainText()).paragraph_format.first_line_indent = Inches(0)
        table.rows[len(table.rows) - 1].cells[4].add_paragraph(ui.t4_ocenka.toPlainText()).paragraph_format.first_line_indent = Inches(0)
        table.rows[len(table.rows) - 1].cells[5].add_paragraph(ui.t4_form.toPlainText()).paragraph_format.first_line_indent = Inches(0)

        document.save(arrWords[arrI[0]])
        ui.t4_nomer.setPlainText('')
        ui.t4_name.setPlainText('')
        ui.t4_trud.setPlainText('')
        ui.t4_ocenka.setPlainText('')
        ui.t4_form.setPlainText('')

    def table5():
        document = Document(arrWords[arrI[0]])

        table = document.tables[8]

        table.add_row()

        arrCount_t5[0] +=1
        table.rows[len(table.rows) - 1].cells[0].add_paragraph(str(arrCount_t5[0])).paragraph_format.first_line_indent = Inches(0)
        table.rows[len(table.rows) - 1].cells[1].add_paragraph(ui.t5_tema.toPlainText()).paragraph_format.first_line_indent = Inches(0)
        table.rows[len(table.rows) - 1].cells[2].add_paragraph(ui.t5_zadanie.toPlainText()).paragraph_format.first_line_indent = Inches(0)
        table.rows[len(table.rows) - 1].cells[3].add_paragraph(ui.t5_form.toPlainText()).paragraph_format.first_line_indent = Inches(0)
        table.rows[len(table.rows) - 1].cells[4].add_paragraph(ui.t5_idk.toPlainText()).paragraph_format.first_line_indent = Inches(0)

        document.save(arrWords[arrI[0]])
        ui.t5_tema.setPlainText('')
        ui.t5_zadanie.setPlainText('')
        ui.t5_form.setPlainText('')
        ui.t5_idk.setPlainText('')


    def table6():
        document = Document(arrWords[arrI[0]])

        table = document.tables[12]

        table.add_row()

        arrHours.append(int(ui.t6_hours.toPlainText()))
        arrCount_t6[0] +=1
        table.rows[len(table.rows) - 1].cells[0].add_paragraph(str(arrCount_t6[0])).paragraph_format.first_line_indent = Inches(0)
        table.rows[len(table.rows) - 1].cells[1].add_paragraph(ui.t6_tema.toPlainText()).paragraph_format.first_line_indent = Inches(0)
        table.rows[len(table.rows) - 1].cells[2].add_paragraph(ui.t6_vid.toPlainText()).paragraph_format.first_line_indent = Inches(0)
        table.rows[len(table.rows) - 1].cells[3].add_paragraph(ui.t6_form.toPlainText()).paragraph_format.first_line_indent = Inches(0)
        table.rows[len(table.rows) - 1].cells[4].add_paragraph(ui.t6_hours.toPlainText()).paragraph_format.first_line_indent = Inches(0)

        document.save(arrWords[arrI[0]])
        ui.t6_tema.setPlainText('')
        ui.t6_vid.setPlainText('')
        ui.t6_form.setPlainText('')
        ui.t6_hours.setPlainText('')

    def table6_count():
        document = Document(arrWords[arrI[0]])
        table = document.tables[12]
        table.add_row()
        table.rows[len(table.rows) - 1].cells[0].merge(table.rows[len(table.rows) - 1].cells[3])

        table.rows[len(table.rows) - 1].cells[0].add_paragraph("Итого часов").paragraph_format.first_line_indent = Inches(0)
        table.rows[len(table.rows) - 1].cells[4].add_paragraph(str(sum(arrHours))).paragraph_format.first_line_indent = Inches(0)
        document.save(arrWords[arrI[0]])

    def table7():
        document = Document(arrWords[arrI[0]])

        table = document.tables[13]

        table.add_row()

        arrCount_t7[0] +=1
        table.rows[len(table.rows) - 1].cells[0].add_paragraph(str(arrCount_t7[0])).paragraph_format.first_line_indent = Inches(0)
        table.rows[len(table.rows) - 1].cells[1].add_paragraph(ui.t7_vid.toPlainText()).paragraph_format.first_line_indent = Inches(0)
        table.rows[len(table.rows) - 1].cells[2].add_paragraph(ui.t7_contr.toPlainText()).paragraph_format.first_line_indent = Inches(0)

        if ui.t7_com.toPlainText() == '':
            table.rows[len(table.rows) - 1].cells[3].merge((table.rows[len(table.rows) - 2].cells[3]))
        else:
            table.rows[len(table.rows) - 1].cells[3].add_paragraph(
                ui.t7_com.toPlainText()).paragraph_format.first_line_indent = Inches(0)

        document.save(arrWords[arrI[0]])
        ui.t7_vid.setPlainText('')
        ui.t7_contr.setPlainText('')

    def page1():

            doc = DocxTemplate(arrWords[arrI[0]])
            context = { 'p1_modul': ui.p1_modul.text(), 'p1_profil': ui.p1_profil.text(), 'p1_ot': ui.p1_ot.text(),
                    'p1_year' : ui.p1_year.text(), 'p1_nomer': ui.p1_nomer.text(), 'p1_protokol' : ui.p1_protokol.text(),
                    'p2_target': ui.p2_target.text(), 'p2_tasks': ui.p2_tasks.toPlainText(), 'p2_modul': ui.p2_modul.text(),
                   'p2_part': ui.p2_part.toPlainText(), 'p2dicisplin': ui.p2_dicisplin.text(), 'p2_predmetov': ui.p2_predmetov.text(),
                   'p22dot2': ui.p2_2dot2.text(), 'p22dot3': ui.p2_2dot3.text(),'p3_trud': ui.p3_trud.toPlainText(),
                    'p3_name': ui.p3_name.toPlainText(), 'p3_forma': ui.p3_forma.toPlainText(),
                   'plainTextEdit_4': ui.plainTextEdit_4.toPlainText(),'p4_45': ui.p4_45.toPlainText(),
                    'p4_main': ui.p4_main.toPlainText(), 'p4_dop': ui.p4_dop.toPlainText(),
                   'p4_period': ui.p4_period.toPlainText(),'pIV_trud': ui.pIV_trud.text(),
                    'pIV_hours': ui.pIV_hours.text(), 'pIV_hoursEx': ui.pIV_hoursEx.text(),
                   'pIV_form': ui.pIV_form.text(),'p5_demka': ui.p5_demka.toPlainText(), 'p6_82': ui.p6_82.toPlainText() }
            doc.render(context)
            doc.save(arrWords[arrI[0]])
            ui.p1_modul.setText('')
            ui.p1_profil.setText('')
            ui.p1_ot.setText('')
            ui.p1_year.setText('')
            ui.p1_nomer.setText('')
            ui.p1_protokol.setText('')
            ui.p2_target.setText('')
            ui.p2_tasks.setPlainText('')
            ui.p2_modul.setText('')
            ui.p2_part.setPlainText('')
            ui.p2_2dot2.setText('')
            ui.p2_2dot3.setText('')
            ui.p3_trud.setPlainText('')
            ui.p3_name.setPlainText('')
            ui.p3_forma.setPlainText('')
            ui.plainTextEdit_4.setPlainText('')
            ui.p4_45.setPlainText('')
            ui.p4_main.setPlainText('')
            ui.p4_dop.setPlainText('')
            ui.p4_period.setPlainText('')
            ui.pIV_trud.setText('')
            ui.pIV_hours.setText('')
            ui.pIV_hoursEx.setText('')
            ui.pIV_form.setText('')
            ui.p5_demka.setPlainText('')
            ui.p6_82.setPlainText('')
            arrI[0] += 1


    arrI = []
    arrI.append(0)
    arrNumber = []
    arrNumber_t4 = []
    arrCount_t4 = []
    arrCount_t4.append(0)
    arrCount_t5 = []
    arrCount_t5.append(0)
    arrCount_t6 = []
    arrCount_t6.append(0)
    arrCount_t7 = []
    arrCount_t7.append(0)
    arrSemestr = []
    arrSeminar = []
    arrSamr = []
    arrHours = []
    t3_sumHours = []
    arrWords = []

    ui.p1_btnSend.clicked.connect(create_word)
    ui.pushButton_5.clicked.connect(table1)
    ui.t2_btnSend.clicked.connect(table2)
    ui.t2_btnToSend.clicked.connect(table2_count)
    ui.t2_btnSend_2.clicked.connect(table3)
    ui.t3_btnSend.clicked.connect(table3_count)
    ui.t4_btnSend.clicked.connect(table4)
    ui.t4_btnSend_2.clicked.connect(table5)
    ui.t6_btnAdd.clicked.connect(table6)
    ui.t6_btnSend.clicked.connect(table6_count)
    ui.t7_btnAdd.clicked.connect(table7)

    ui.p6_btnSend.clicked.connect(page1)
    sys.exit(app.exec_())

